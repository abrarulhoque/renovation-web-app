from docx import Document
import os
import shutil
import re
import logging


def modify_template(base_dir=None):
    """
    Create a new template based on the original Word template with placeholders for the renovation quote.
    This preserves the original formatting while adding the template variables.

    Args:
        base_dir: Optional base directory path. If provided, paths will be absolute.
    """
    logger = logging.getLogger(__name__)

    # Use absolute paths if base_dir is provided
    if base_dir:
        # Path to the original template
        original_template = os.path.join(
            base_dir, "templates", "word", "renovation_quote.docx"
        )
        # Path for the new template
        new_template = os.path.join(
            base_dir, "templates", "word", "renovation_quote_template.docx"
        )
    else:
        # Path to the original template (relative paths)
        original_template = os.path.join("templates", "word", "renovation_quote.docx")
        # Path for the new template
        new_template = os.path.join(
            "templates", "word", "renovation_quote_template.docx"
        )

    logger.debug(f"Original template path: {original_template}")
    logger.debug(f"New template path: {new_template}")

    # Check if original template exists
    if not os.path.exists(original_template):
        logger.error(f"Original template not found at: {original_template}")
        logger.error(f"Current directory: {os.getcwd()}")
        if base_dir:
            logger.error(f"Base directory: {base_dir}")
        raise FileNotFoundError(f"Original template not found: {original_template}")

    # Create a copy of the original template
    shutil.copy2(original_template, new_template)
    print(f"Template copied to {new_template}")

    # Load the document
    doc = Document(new_template)

    # Regular expression to find placeholders in curly braces
    placeholder_pattern = re.compile(r"\{([^{}]+)\}")

    # Track all found placeholders
    found_placeholders = set()
    replacements_made = 0

    # Process all tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Find all placeholders in the paragraph text
                    matches = placeholder_pattern.findall(paragraph.text)

                    if matches:
                        # Replace each placeholder with Jinja2 format
                        text = paragraph.text
                        for placeholder in matches:
                            original = f"{{{placeholder}}}"
                            jinja_format = f"{{{{{placeholder}}}}}"
                            text = text.replace(original, jinja_format)
                            found_placeholders.add(placeholder)
                            replacements_made += 1

                        paragraph.text = text

    # Process all paragraphs in the document (outside tables)
    for paragraph in doc.paragraphs:
        # Find all placeholders in the paragraph text
        matches = placeholder_pattern.findall(paragraph.text)

        if matches:
            # Replace each placeholder with Jinja2 format
            text = paragraph.text
            for placeholder in matches:
                original = f"{{{placeholder}}}"
                jinja_format = f"{{{{{placeholder}}}}}"
                text = text.replace(original, jinja_format)
                found_placeholders.add(placeholder)
                replacements_made += 1

            paragraph.text = text

    # Save the modified template
    doc.save(new_template)

    if replacements_made > 0:
        print(
            f"Template modified successfully with {replacements_made} placeholders at {new_template}"
        )
        print(f"Found placeholders: {', '.join(sorted(found_placeholders))}")
    else:
        print("Warning: No placeholders were found in the template.")

    return new_template


if __name__ == "__main__":
    modify_template()

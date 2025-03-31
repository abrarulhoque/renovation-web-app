from docx import Document
import os


def create_template():
    """
    Create a Word template with variables for the renovation quote.
    """
    # Create a new document
    doc = Document()

    # Add a title
    doc.add_heading("Bathroom Renovation Quote", 0)

    # Add a section for personal details
    doc.add_heading("Personal Details", level=1)

    # Create a table for personal details
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    # Add headers and variable placeholders
    cells = table.rows[0].cells
    cells[0].text = "Name:"
    cells[1].text = "{{name}}"

    cells = table.rows[1].cells
    cells[0].text = "Email:"
    cells[1].text = "{{ePost}}"

    # Add another table for contact details
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    cells = table.rows[0].cells
    cells[0].text = "Phone:"
    cells[1].text = "{{telefon}}"

    cells = table.rows[1].cells
    cells[0].text = "Address:"
    cells[1].text = "{{address}}"

    # Ensure the directory exists
    os.makedirs("templates/word/temp", exist_ok=True)

    # Save the document
    doc.save("templates/word/temp/renovation_quote_template.docx")

    return "templates/word/temp/renovation_quote_template.docx"


if __name__ == "__main__":
    create_template()
